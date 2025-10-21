from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from openai import OpenAI
from io import BytesIO
import re
import traceback
import os



# ------- Word (python-docx) -------
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# ------- PDF (reportlab) -------
import os
import tempfile
from io import BytesIO
import subprocess
from docx import Document
import platform

from datetime import datetime


app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

@app.route("/", methods=["GET"])
def home():
    return "Resume Automation API is live 🚀. Use /submit with POST."

# ---- Section detection ----
SECTION_TITLES = {
    "professional summary",
    "summary",
    "technical skills",
    "skills",
    "professional experience",
    "education",
    "certifications",
    "experience",
    "work experience",
    "work history",
    "projects",
    "additional qualifications",
    "additional information",
    "references",
}

def clean_markdown(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
    text = text.replace("`", "")
    text = re.sub(r"^\s{0,3}#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    text = text.replace("**", "").replace("*", "").replace("_", "")
    text = re.sub(r"^\s*[•\-–]\s*", "- ", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def is_contact_line(line: str) -> bool:
    if not line:
        return False
    l = line.lower()
    return ("email" in l or "@" in l or "phone" in l or re.search(r"\b\d{10}\b", l) or re.search(r"\+\d", l))

def is_section_title(line: str) -> bool:
    if not line:
        return False
    raw = line.strip().rstrip(":")
    return raw.lower() in SECTION_TITLES

def add_horizontal_rule(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ---- Word building helpers ----
def add_candidate_name(doc, lines, idx):
    if idx < len(lines):
        name_para = doc.add_paragraph(lines[idx])
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = name_para.runs[0]
        run.bold = True
        run.font.size = Pt(20)
        idx += 1
    return idx

def add_contact_info(doc, lines, idx):
    contact_email, contact_phone, contact_location = "", "", ""
    while idx < len(lines) and is_contact_line(lines[idx]):
        line = lines[idx]
        email_match = re.search(r"[\w\.-]+@[\w\.-]+", line)
        if email_match:
            contact_email = email_match.group(0)
        phone_match = re.search(r"(\+?\d[\d\s\-]{8,}\d)", line)
        if phone_match:
            contact_phone = phone_match.group(0).strip()
        loc_match = re.search(r"Location\s*[:\-]?\s*(.*)", line, re.IGNORECASE)
        if loc_match:
            contact_location = loc_match.group(1).strip()
        idx += 1

    if contact_email or contact_phone or contact_location:
        pieces = []
        if contact_email:
            pieces.append(f"Email: {contact_email}")
        if contact_phone:
            pieces.append(f"Mobile: {contact_phone}")
        if contact_location:
            pieces.append(f"Location: {contact_location}")
        contact_line = "  |  ".join(pieces)
        contact_para = doc.add_paragraph(contact_line)
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact_para.runs[0].font.size = Pt(11)

    return idx


def add_section_title(doc, title, idx):
    p = doc.add_paragraph(title.upper().rstrip(":"))
    p.paragraph_format.space_before = Pt(12)   # add spacing above the title
    p.paragraph_format.space_after = Pt(4)     # small space below (optional)

    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(12)                       # set font size to 12

    add_horizontal_rule(p)
    return idx + 1


def add_skills_section(doc, lines, idx):
    idx = add_section_title(doc, lines[idx], idx)
    category = None
    skills = []

    while idx < len(lines) and not is_section_title(lines[idx]):
        line = lines[idx].strip()

        if not line:
            idx += 1
            continue

        # Case 1: Inline format (comma-separated skills, no "-")
        if category and not line.startswith("-") and "," in line:
            skills = [s.strip() for s in line.split(",") if s.strip()]
            p = doc.add_paragraph()
            r1 = p.add_run(category + ": ")
            r1.bold = True
            r2 = p.add_run(", ".join(skills))
            category, skills = None, []  # reset after flush

        # Case 2: New category line
        elif not line.startswith("-"):
            if category and skills:
                p = doc.add_paragraph()
                r1 = p.add_run(category + ": ")
                r1.bold = True
                r2 = p.add_run(", ".join(skills))
            category = line
            skills = []

        # Case 3: Bulleted skill
        else:
            skills.append(line.lstrip("- ").strip())

        idx += 1

    # flush last category
    if category and skills:
        p = doc.add_paragraph()
        r1 = p.add_run(category + ": ")
        r1.bold = True
        r2 = p.add_run(", ".join(skills))

    return idx


    # flush last category
    if category and skills:
        p = doc.add_paragraph()
        r1 = p.add_run(category + ": ")
        r1.bold = True
        r2 = p.add_run(", ".join(skills))

    return idx


def add_experience_section(doc, lines, idx):
    idx = add_section_title(doc, lines[idx], idx)
    company_seen = False  # track first company
    while idx < len(lines) and not is_section_title(lines[idx]):
        line = lines[idx]

        # ✅ Company – Location OR Role – Dates
        if " – " in line and ":" not in line:
            if " to " in line:  # ✅ Role line
                p = doc.add_paragraph(line)
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(10)
                # p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                # ❌ No space above for role lines
            else:  # ✅ Company line
                p = doc.add_paragraph(line)
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(11)

                # ✅ Add space above only for companies after the first one
                if company_seen:
                    p.paragraph_format.space_before = Pt(10)
                company_seen = True

        elif " – " in line and ":" in line:  # job + bullet description
            job_title, rest = line.split(":", 1)
            p = doc.add_paragraph(job_title.strip())
            p.runs[0].bold = True
            parts = re.split(r'\.\s+|,\s+', rest)
            for part in parts:
                if part.strip():
                    bullet_para = doc.add_paragraph(part.strip(), style="List Bullet")
                    bullet_para.paragraph_format.left_indent = Inches(0.25)

        elif line.startswith("Technologies Used"):
            heading, _, techs = line.partition(":")
            p = doc.add_paragraph()
            r1 = p.add_run(heading.strip() + ": ")
            r1.bold = True
            p.add_run(techs.strip())
            # ✅ Only spacing below (no space above)
            p.paragraph_format.space_after = Pt(10)

        elif line.startswith("- "):  # bullets with -
            bullet_para = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            bullet_para.paragraph_format.left_indent = Inches(0.25)

        else:
            doc.add_paragraph(line)

        idx += 1
    return idx


def add_certifications_section(doc, lines, idx):
    idx = add_section_title(doc, lines[idx], idx)
    while idx < len(lines) and not is_section_title(lines[idx]):
        line = lines[idx].lstrip("- ").strip()   # ✅ remove leading "-"
        if line:
            doc.add_paragraph("• " + line)
        idx += 1
    return idx


def add_education_section(doc, lines, idx):
    idx = add_section_title(doc, lines[idx], idx)
    while idx < len(lines) and not is_section_title(lines[idx]):
        doc.add_paragraph(lines[idx])
        idx += 1
    return idx

def add_summary_section(doc, lines, idx):
    idx = add_section_title(doc, lines[idx], idx)  # add "PROFESSIONAL SUMMARY"
    
    while idx < len(lines) and not is_section_title(lines[idx]):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue

        # Always force bullet format (whether line starts with "- " or not)
        if line.startswith("- "):
            text = line[2:].strip()
        else:
            text = line

        bullet_para = doc.add_paragraph(text, style="List Bullet")
        bullet_para.paragraph_format.left_indent = Inches(0.25)
        idx += 1

    return idx



def extract_total_experience(candidate_info: str) -> str:
    # Normalize dashes
    candidate_info = candidate_info.replace("–", "-").replace("—", "-")

    # Extract all duration lines
    duration_lines = re.findall(r"Duration:\s*(.+)", candidate_info, re.IGNORECASE)

    total_months = 0
    today = datetime.today()

    for line in duration_lines:
        # Split into start and end
        parts = [p.strip() for p in line.split("-")]
        if len(parts) != 2:
            continue
        
        start_str, end_str = parts

        # Parse start date
        try:
            start_date = datetime.strptime(start_str, "%b %Y")
        except ValueError:
            start_date = datetime.strptime(start_str, "%B %Y")

        # Handle "Present"
        if "present" in end_str.lower():
            end_date = today
        else:
            try:
                end_date = datetime.strptime(end_str, "%b %Y")
            except ValueError:
                end_date = datetime.strptime(end_str, "%B %Y")

        # Calculate duration in months
        months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
        total_months += months

        # Print each duration separately
        years, m = divmod(months, 12)
        # print(f"{start_str} – {end_str}: {years} years {m} months")

    # Convert total months into years+months
    total_years, total_m = divmod(total_months, 12)
    print(f"Total Experience: {total_years} years {total_m} months")
    return f"Total Experience: {total_years} years {total_m} months"


# ---- Main Word generator ----
def create_resume_word(content: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    # 👉 Add this block
    para_format = style.paragraph_format
    para_format.space_after = Pt(0)
    para_format.space_before = Pt(0)
    para_format.line_spacing = 1
    para_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    lines = [ln.strip("• ").strip() for ln in content.splitlines() if ln and str(ln).strip()]
    idx = 0

    # Candidate Name
    idx = add_candidate_name(doc, lines, idx)

    # Contact Info
    idx = add_contact_info(doc, lines, idx)

    # Sections
    while idx < len(lines):
        if is_section_title(lines[idx]):
            section_key = lines[idx].strip().rstrip(":").lower()
            if section_key in ("professional summary", "summary"):
                idx = add_summary_section(doc, lines, idx)   # ✅ FIXED
            elif section_key in ("skills", "technical skills"):
                idx = add_skills_section(doc, lines, idx)
            elif section_key in ("work experience", "professional experience"):
                idx = add_experience_section(doc, lines, idx)
            elif section_key == "certifications":
                idx = add_certifications_section(doc, lines, idx)
            elif section_key == "education":
                idx = add_education_section(doc, lines, idx)
            else:
                idx = add_section_title(doc, lines[idx], idx)
        else:
            idx += 1

    return doc



def create_resume_pdf(resume_text: str) -> BytesIO:
    # Step 1: Create Word doc
    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc = create_resume_word(resume_text)
    doc.save(tmp_docx.name)

    # Step 2: Prepare PDF path
    tmp_pdf_path = os.path.splitext(tmp_docx.name)[0] + ".pdf"

    # Step 3: Determine LibreOffice executable path
    system = platform.system()
    if system == "Windows":
        # Change this path if LibreOffice installed elsewhere
        soffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
    else:
        soffice_path = "libreoffice"  # Linux / Mac assumes in PATH

    # Step 4: Convert DOCX -> PDF
    try:
        subprocess.run([
            soffice_path,
            "--headless",
            "--convert-to", "pdf",
            tmp_docx.name,
            "--outdir", os.path.dirname(tmp_pdf_path)
        ], check=True)
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"LibreOffice PDF conversion failed: {e}")
    except FileNotFoundError:
        raise RuntimeError(f"LibreOffice not found at {soffice_path}. Install it or update the path.")

    # Step 5: Read PDF
    with open(tmp_pdf_path, "rb") as f:
        pdf_bytes = BytesIO(f.read())

    # Step 6: Cleanup
    try:
        os.remove(tmp_docx.name)
        os.remove(tmp_pdf_path)
    except OSError:
        pass

    pdf_bytes.seek(0)
    return pdf_bytes


# ---- API endpoint ----
@app.route("/submit", methods=["POST"])
def submit():
    try:
        data = request.get_json(force=True, silent=False)
    except Exception:
        return jsonify({"message": "Invalid JSON"}), 400

    job_desc = (data or {}).get("job_desc", "").strip()
    candidate_info = (data or {}).get("candidate_info", "").strip()
    file_type = (data or {}).get("file_type", "word").strip().lower()

    work_exp_str = extract_total_experience(candidate_info)

    if not job_desc or not candidate_info:
        return jsonify({"message": "Missing required fields"}), 400

    try:
        client = OpenAI(api_key=OPENAI_API_KEY)
        prompt = f"""
        You are a professional resume writer. Using the Job Description and Candidate Information provided below, generate a clean, ATS-optimized resume that strictly follows the section order and formatting rules listed here:

        ⚠️ IMPORTANT: Output must contain the **resume only** — do not include explanations, disclaimers, notes, or extra text outside of the resume.

        SECTION ORDER:

        1. **PROFESSIONAL SUMMARY** – Include **6 to 8 bullet points**.  
            - The **first bullet point must always mention the candidate's total years of professional experience**
            WORK Experience: {work_exp_str}  
            - Represent the total as "X+ years of experience" (e.g., 5+ years, 6+ years), based **strictly on the earliest start date and the latest end year found in the CANDIDATE INFORMATION**, ignoring any "Present" or current date mentions.  
            - Do not infer, estimate, or change the experience from the Job Description or any other source.  
            - The remaining 2 liner detailed bullet points (5–7) must highlight key skills, achievements, career highlights, and qualifications aligned with the Job Description.  
            - Each bullet must start with "- ".  

        2. **SKILLS** – Based on the Job Description and Candidate Information:

            1. Identify the **most relevant role/position** (e.g., .NET Developer, Java Backend Engineer, Salesforce Developer, Data Engineer, DevOps Engineer).
            2. Create a **resume-ready Skills section** with **10–12 subsections**, tailored to that role and the JD.

            ⚠️ RULES:
            - Subsections must be **category-based** and recruiter-friendly (e.g., Programming Languages, Frameworks & Libraries, Databases, Cloud Platforms, DevOps & CI/CD, Testing & QA, Security & Compliance, Monitoring & Observability, Collaboration Tools).
            - Use concise, ATS-optimized, professional wording for subsection titles.
            - Fill each subsection with **8–20 related technologies/tools**, directly matching the JD and candidate info.
            - Where possible, **expand categories with specific services or tools** (e.g., list AWS services like EC2, S3, Glue, Lambda, CloudWatch — not just "AWS").
            - Always mirror exact JD keywords (e.g., if JD says “GCP, Spark, BigQuery, Kafka” → those must appear under correct categories).
            - Include versions where impactful (e.g., Java 11/17, .NET 6/7, Spring Boot 3.x, Hadoop 3.x).
            - Do not invent irrelevant categories or mix unrelated technologies into the wrong subsection.
            - Always include these **mandatory baseline categories**, even if not explicitly in the JD:
                - Programming Languages  
                - Operating Systems  
                - Cloud Platforms
                - DevOps & CI/CD Tools  
                - Development Tools                   

            Example subsections (adjust dynamically per JD):  
            - Programming Languages  
            - Frameworks & Libraries  
            - Databases & Data Warehousing  
            - Big Data & Streaming  
            - Cloud Platforms  
            - DevOps & CI/CD Tools  
            - Testing & QA  
            - Security & Compliance  
            - Monitoring & Observability  
            - Collaboration Tools  
            - Documentation Tools  
            - Operating Systems  

            ⚠️ Ensure each subsection is **fully loaded with at least 8 skills** and contains **16–20 skills where possible**.
            ⚠️ All technologies listed here must also appear in the **Technologies Used** lines under the WORK EXPERIENCE section.



        3. **CERTIFICATIONS**

        4. **EDUCATION**
        
        ⚠️ IMPORTANT: Do NOT generate the WORK EXPERIENCE section. It will be added separately.

        FORMATTING RULES:
        - Display the candidate’s **Name** at the top.
        - Center **Email**, **Phone Number**, and **Candidate Location** on the same line directly below the name, using the format:  
        Email: | Mobile: | Location:
        - Use 0.5-inch page margins.
        - Add a tab space before each bullet point.
        - Do not use markdown or bullet characters like "-", "*", or "•".
        - The **SKILLS** section must always follow the defined categories above—never as a plain list.
        - Always ensure the final resume spans at least 2 full pages of Word or PDF output.

        JOB DESCRIPTION:
        {job_desc}

        CANDIDATE INFORMATION:
        {candidate_info}
        """
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You write polished, ATS-friendly resumes."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.6,
        )
        raw_resume = resp.choices[0].message.content or ""


         # ----------- SECOND CALL (ONLY Work Experience Section) -----------
        exp_prompt = f"""
        Generate ONLY the WORK EXPERIENCE section for this resume.

        3. **WORK EXPERIENCE** – Merge **Work History** and **Work Experience** into a unified section. For each job role:
            - ⚠️ IMPORTANT: Use WORK EXPERIENCE from the CANDIDATE INFORMATION only
            - Include the Job Title, Company Name (bold), Job Location, and timeline using the format:
                [Company Name] – [Job Location]  
                [Job Title] – [Start Month Year] to [End Month Year]

        - Add 10 to 15 high-impact bullet points per role. Each bullet point must:
        - Each bullet point must be exactly 2 lines long, with rich and specific details — including technologies used, metrics, project outcomes, team collaboration, challenges faced, and business impact.
        - "When generating points for each company, first identify the industry it operates in, and then tailor the points to be relevant to that specific industry projects.
        - Start with a strong action verb (e.g., Spearheaded, Engineered, Optimized, Automated, Delivered).
        - Focus on achievements, measurable outcomes, and business value rather than just responsibilities.
        - Include quantifiable results wherever possible (e.g., improved ETL performance by 35%, reduced deployment time by 40%, cut costs by 20% annually).
        - Highlight leadership, innovation, automation, and cross-functional collaboration.
        - Showcase modern practices (e.g., Cloud Migration, DevOps, CI/CD automation, Data Engineering, AI/ML, Security, Scalability).
        - Be specific, technical, and results-driven — not generic.
        

        - ⚠️ Validate technology usage against the job timeline:
        - ONLY include technologies, tools, frameworks, or platforms that were **publicly available and in practical use** during the given employment period.
        - Example: Do NOT include Generative AI, Azure OpenAI, MS Fabric, or other technologies launched post-2021 in roles dated 2020 or earlier.
        - Ensure all technologies and practices mentioned are **realistically applicable** based on release year and industry adoption timeline.

        - Total bullet points should follow this logic:
        - For 1 company: 15 to 20 bullet points.
        - For 2 companies: 15 to 20 bullet points each (total: 30-40 points).
        - For 3 companies: 10 to 15 bullet points each (total: 30-45 points).
        - For 4 companies: 10 to 15 bullet points each (total: 40-60 points).
        - For 5 companies: 10 to 15 bullet points each (total: 60-70 points).
        - For 6 companies: 10 to 15 bullet points each (total: 70-80 points).
        - For 7 companies: 10 to 15 bullet points each (total: 70-80 points).

        - No filler or repetition: Each bullet point must offer unique, concrete contributions or achievements.

        - Write in professional resume tone, use strong action verbs, and focus on clarity, impact, and relevance to technical or engineering roles.

        - End each job section with the line:  
        Technologies Used: tech1, tech2, ..., tech15  
            ⚠️ Ensure each role includes 10 to 15 technologies mapped directly from the SKILLS section.  
            ⚠️ Across all roles, the union of technologies must comprehensively cover the entire SKILLS section.

        JOB DESCRIPTION:
        {job_desc}

        CANDIDATE INFORMATION:
        {candidate_info}
        """

        resp_exp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You write only the Work Experience section for ATS resumes."},
                {"role": "user", "content": exp_prompt},
            ],
            temperature=0.6,
        )
        exp_text = resp_exp.choices[0].message.content or ""
        
        # ✅ MERGE: Append Work Experience at the end
        main_content = clean_markdown(raw_resume).strip()
        work_exp_content = clean_markdown(exp_text).strip()
        
        # Ensure work experience has proper title
        if work_exp_content and not work_exp_content.upper().startswith("WORK EXPERIENCE"):
            work_exp_content = "WORK EXPERIENCE\n" + work_exp_content
        
        # Append Work Experience at the end (after Certifications and Education)
        resume_text = main_content + "\n\n" + work_exp_content
        
    except Exception as e:
        traceback.print_exc()
        return jsonify({"message": f"OpenAI error: {e}"}), 500

    if not resume_text:
        return jsonify({"message": "Resume generation failed: Empty response from AI"}), 500
    
    # ✅ Extract candidate name (first line of resume_text)
    candidate_name = resume_text.splitlines()[0].strip()
    safe_name = re.sub(r'[^A-Za-z0-9]+', '_', candidate_name)  # replace spaces & symbols

    try:
        if file_type == "word":
            buffer = BytesIO()
            doc = create_resume_word(resume_text)
            doc.save(buffer)
            buffer.seek(0)
            return send_file(
                buffer,
                as_attachment=True,
                download_name = safe_name + "_resume.docx",   # ✅ dynamic name
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif file_type == "pdf":
            buffer = create_resume_pdf(resume_text)
            return send_file(buffer, as_attachment=True, download_name="resume.pdf", mimetype="application/pdf")
        else:
            return jsonify({"message": "Invalid file_type"}), 400
    except Exception as e:
        traceback.print_exc()
        return jsonify({"message": f"File generation error: {e}"}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True) # production
    # app.run(host="127.0.0.1", port=5000, debug=True) # local testing